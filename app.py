import streamlit as st
import requests
import json
from io import BytesIO
from docx import Document

# Configurar la página
title = "Generador de correos profesionales en español"
st.set_page_config(page_title=title)
st.title(title)

# Inicializar el contador de visitas
if 'visit_count' not in st.session_state:
    st.session_state['visit_count'] = 100

st.session_state['visit_count'] += 1

# Mostrar el contador de visitas y los enlaces en la barra lateral
st.sidebar.write(f"Visitas a la aplicación: {st.session_state['visit_count']}")
st.sidebar.write("**¿Desea revisión profesional de textos en 24 horas? [Visite hablemosbien.org](https://hablemosbien.org)**")
st.sidebar.write("**Su anuncio [aquí](https://hablemosbien.org/anuncio.htm)**")

# Descripción de la aplicación
st.write("""
Esta aplicación genera correos profesionales en español basados en el mensaje del usuario y la información del destinatario. 
Proporciona un borrador con tono profesional, adaptado a las normas culturales del ámbito hispanohablante.
""")

# Entrada del usuario
user_message = st.text_area("Ingresa el mensaje breve que deseas transmitir:", "")
recipient_info = st.text_input("Ingresa información sobre el destinatario (por ejemplo, su nombre y relación contigo):", "")

tone_options = ["Formal", "Cordial", "Informal"]
tone = st.selectbox("Selecciona el tono del correo:", tone_options)

# Botón para generar el correo
if st.button("Generar correo"):
    if not user_message or not recipient_info:
        st.warning("Por favor, completa todos los campos antes de continuar.")
    else:
        # Preparar el prompt para la API
        prompt = f"""
        Contexto: Eres un experto en redacción de correos profesionales en español. Tu tarea es crear un borrador de correo basado en el siguiente mensaje:
        
        Mensaje del usuario: {user_message}
        Información del destinatario: {recipient_info}
        Tono: {tone}
        
        El correo debe estar estructurado de la siguiente manera:
        - Saludo apropiado
        - Introducción con el propósito del mensaje
        - Cuerpo principal con el contenido detallado
        - Cierre con resumen y llamada a la acción
        - Despedida profesional y firma
        """

        # Hacer la solicitud a la API de Together
        api_url = "https://api.together.xyz/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {st.secrets['TOGETHER_API_KEY']}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "deepseek-ai/DeepSeek-R1",
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 1000,
            "temperature": 0.7,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 1,
            "stop": ["<|end▁of▁sentence|>"]
        }

        response = requests.post(api_url, headers=headers, data=json.dumps(payload))

        if response.status_code == 200:
            response_data = response.json()
            email_content = response_data['choices'][0]['message']['content']

            # Mostrar el correo generado
            st.subheader("Correo generado")
            st.write(email_content)

            # Opción para descargar el correo en formato DOCX
            doc = Document()
            doc.add_heading("Correo profesional generado", level=1)
            doc.add_paragraph(email_content)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="Descargar correo en DOCX",
                data=buffer,
                file_name="correo_profesional.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("Hubo un error al generar el correo. Por favor, intenta nuevamente.")
